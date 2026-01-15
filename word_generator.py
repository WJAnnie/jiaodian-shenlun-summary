"""Word格子纸生成模块"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_cell_border(cell, color="FF0000"):
    """设置单元格红色边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcPr.append(border)


def create_shenlun_doc(text, date_str, output_dir):
    """创建申论格子纸Word文档"""
    # 清理文本
    clean_text = text.replace('【标题】', '').replace('【正文】', '')
    clean_text = clean_text.replace('\n\n', '\n').strip()

    # 参数设置
    cols = 25  # 每行格子数
    rows = 20  # 每页行数
    chars_per_page = cols * rows

    # 分页
    pages = []
    chars = list(clean_text.replace('\n', ''))
    for i in range(0, len(chars), chars_per_page):
        pages.append(chars[i:i+chars_per_page])

    doc_paths = []

    for page_idx, page_chars in enumerate(pages):
        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(f"{date_str}焦点访谈申论总结")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.name = 'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

        # 创建表格
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 填充文字
        char_idx = 0
        for row in table.rows:
            row.height = Cm(0.8)
            for cell in row.cells:
                cell.width = Cm(0.7)
                set_cell_border(cell)
                if char_idx < len(page_chars):
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(page_chars[char_idx])
                    run.font.size = Pt(12)
                    run.font.name = 'KaiTi'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'KaiTi')
                    char_idx += 1

        # 保存
        filename = f"申论_{date_str}_{page_idx+1}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        doc_paths.append(filepath)
        print(f"已生成: {filename}")

    return doc_paths
